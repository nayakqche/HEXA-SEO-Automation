"""
Upload container — a persistent bag of user-supplied source material that gets
folded into the PRIMARY grounding context for a run.

The user drops in files (txt / json / csv / pdf / doc(x) / xlsx / images / mp4)
or pastes raw text. Each is parsed into one of a few kinds:

  • text   — plain text extracted from the file (txt, json, pdf, docx paragraphs)
  • table  — structured rows extracted from csv / xlsx / docx tables
  • image  — a real photo / chart / graph / table screenshot the blog must keep
  • video  — stored as-is; content isn't machine-readable, so it's noted only
  • other  — stored, noted, not embedded

Everything lives under `uploads/` with a single `manifest.json` describing the
container. `resource_context()` turns the container into (primary_text,
embeddable_images, embeddable_tables) for the pipeline. `clear()` empties it.

Nothing here talks to the network or the LLM — it's pure local parsing so a
malformed upload can never break a run; it just gets skipped with a note.
"""

from __future__ import annotations

import csv
import io
import json
import re
import secrets
import shutil
from pathlib import Path

UPLOAD_DIR = Path("uploads")
MANIFEST = UPLOAD_DIR / "manifest.json"

# Per-resource text budget so one giant PDF can't blow up the prompt.
_TEXT_BUDGET = 12000
_TABLE_ROW_LIMIT = 60

_IMAGE_EXT = {".png", ".jpg", ".jpeg", ".gif", ".webp"}
_VIDEO_EXT = {".mp4", ".mov", ".webm", ".mkv"}


# ── Manifest I/O ────────────────────────────────────────────────────────────

def _load_manifest() -> dict:
    if not MANIFEST.exists():
        return {"resources": []}
    try:
        return json.loads(MANIFEST.read_text(encoding="utf-8"))
    except (json.JSONDecodeError, OSError):
        return {"resources": []}


def _save_manifest(manifest: dict) -> None:
    UPLOAD_DIR.mkdir(exist_ok=True)
    MANIFEST.write_text(json.dumps(manifest, indent=2, ensure_ascii=False),
                        encoding="utf-8")


def _new_id() -> str:
    return "up" + secrets.token_hex(4)


def _safe_name(name: str) -> str:
    name = re.sub(r"[^\w.\-]+", "_", (name or "file").strip())
    return name[:80] or "file"


# ── Parsers ─────────────────────────────────────────────────────────────────

def _rows_to_text(rows: list[list[str]]) -> str:
    """Render extracted rows as a compact pipe table for the LLM context."""
    trimmed = [[("" if c is None else str(c)).strip() for c in r] for r in rows]
    trimmed = [r for r in trimmed if any(cell for cell in r)]
    return "\n".join(" | ".join(r) for r in trimmed[:_TABLE_ROW_LIMIT])


def _parse_csv(data: bytes) -> dict:
    text = data.decode("utf-8-sig", errors="replace")
    rows = [r for r in csv.reader(text.splitlines()) if any(c.strip() for c in r)]
    return {"kind": "table", "rows": rows[:_TABLE_ROW_LIMIT],
            "text": _rows_to_text(rows), "note": f"{len(rows)} row(s)"}


def _parse_xlsx(data: bytes) -> dict:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    rows: list[list[str]] = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            rows.append(["" if v is None else str(v) for v in row])
            if len(rows) >= _TABLE_ROW_LIMIT:
                break
    wb.close()
    rows = [r for r in rows if any(c.strip() for c in r)]
    return {"kind": "table", "rows": rows, "text": _rows_to_text(rows),
            "note": f"{len(rows)} row(s)"}


def _parse_docx(data: bytes) -> dict:
    from docx import Document

    doc = Document(io.BytesIO(data))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables: list[list[list[str]]] = []
    for t in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        rows = [r for r in rows if any(cell for cell in r)]
        if rows:
            tables.append(rows[:_TABLE_ROW_LIMIT])
    text = "\n".join(paras)
    for rows in tables:
        text += "\n\nTABLE:\n" + _rows_to_text(rows)
    # If the doc is mostly one clean table, expose it as a table resource too.
    kind = "table" if (tables and len(paras) <= 2) else "text"
    first_rows = tables[0] if tables else []
    return {"kind": kind, "text": text[:_TEXT_BUDGET], "rows": first_rows,
            "note": f"{len(paras)} paragraph(s), {len(tables)} table(s)"}


def _parse_pdf(data: bytes) -> dict:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    chunks = []
    for page in reader.pages:
        try:
            chunks.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 — never let one bad page kill parsing
            continue
    text = re.sub(r"\n{3,}", "\n\n", "\n".join(chunks)).strip()
    return {"kind": "text", "text": text[:_TEXT_BUDGET],
            "note": f"{len(reader.pages)} page(s)"}


def _parse_text(data: bytes, *, as_json: bool) -> dict:
    text = data.decode("utf-8-sig", errors="replace")
    if as_json:
        try:
            text = json.dumps(json.loads(text), indent=2, ensure_ascii=False)
        except json.JSONDecodeError:
            pass
    return {"kind": "text", "text": text[:_TEXT_BUDGET], "note": "text"}


# ── Public API ──────────────────────────────────────────────────────────────

def add_file(filename: str, data: bytes, description: str = "") -> dict:
    """Parse one upload, store it, append it to the manifest, return its record."""
    UPLOAD_DIR.mkdir(exist_ok=True)
    rid = _new_id()
    safe = _safe_name(filename)
    ext = Path(safe).suffix.lower()
    stored_as = f"{rid}_{safe}"
    (UPLOAD_DIR / stored_as).write_bytes(data)

    record: dict = {
        "id": rid, "name": safe, "stored_as": stored_as,
        "description": (description or "").strip(),
    }

    try:
        if ext in _IMAGE_EXT:
            record.update(kind="image", note="image (kept for the blog)")
        elif ext in _VIDEO_EXT:
            record.update(kind="video", note="video (stored, not embedded in text formats)")
        elif ext == ".csv":
            record.update(_parse_csv(data))
        elif ext == ".xlsx":
            record.update(_parse_xlsx(data))
        elif ext in (".docx", ".doc"):
            record.update(_parse_docx(data))
        elif ext == ".pdf":
            record.update(_parse_pdf(data))
        elif ext == ".json":
            record.update(_parse_text(data, as_json=True))
        elif ext in (".txt", ".md", ""):
            record.update(_parse_text(data, as_json=False))
        else:
            record.update(kind="other", note=f"stored ({ext or 'no extension'})")
    except Exception as exc:  # noqa: BLE001 — surface the problem, keep the file
        record.update(kind="other", note=f"couldn't parse: {exc}")

    manifest = _load_manifest()
    manifest["resources"].append(record)
    _save_manifest(manifest)
    return record


def add_pasted_text(text: str, name: str = "Pasted note") -> dict:
    return add_file(f"{_safe_name(name)}.txt", (text or "").encode("utf-8"))


def list_resources() -> list[dict]:
    """Public-facing resource list (no bulky extracted text)."""
    out = []
    for r in _load_manifest()["resources"]:
        out.append({
            "id": r["id"], "name": r["name"], "kind": r.get("kind", "other"),
            "note": r.get("note", ""), "description": r.get("description", ""),
        })
    return out


def remove(resource_id: str) -> bool:
    manifest = _load_manifest()
    kept, removed = [], None
    for r in manifest["resources"]:
        if r["id"] == resource_id:
            removed = r
        else:
            kept.append(r)
    if removed is None:
        return False
    _quiet_unlink(UPLOAD_DIR / removed.get("stored_as", ""))
    manifest["resources"] = kept
    _save_manifest(manifest)
    return True


def clear() -> None:
    """Empty the whole container so a new run can start with fresh resources."""
    if UPLOAD_DIR.exists():
        shutil.rmtree(UPLOAD_DIR, ignore_errors=True)


def _quiet_unlink(path: Path) -> None:
    try:
        path.unlink()
    except OSError:
        pass


def resource_context() -> tuple[str, list[dict], list[dict]]:
    """
    Fold the container into pipeline inputs:

      returns (primary_text, images, tables)

        primary_text — extracted text/tables, formatted for the grounding prompt
        images       — [{id, name, stored_as, description}] to embed as image blocks
        tables       — [{id, name, rows, description}] to embed as table blocks
    """
    manifest = _load_manifest()
    text_parts: list[str] = []
    images: list[dict] = []
    tables: list[dict] = []

    for r in manifest["resources"]:
        kind = r.get("kind", "other")
        label = r["name"] + (f" — {r['description']}" if r.get("description") else "")
        if kind == "image":
            images.append({"id": r["id"], "name": r["name"],
                           "stored_as": r["stored_as"],
                           "description": r.get("description", "")})
        elif kind == "table":
            if r.get("rows"):
                tables.append({"id": r["id"], "name": r["name"],
                               "rows": r["rows"],
                               "description": r.get("description", "")})
            if r.get("text"):
                text_parts.append(f"\n===== UPLOADED TABLE: {label} =====\n{r['text']}")
        elif kind in ("text",):
            if r.get("text"):
                text_parts.append(f"\n===== UPLOADED FILE: {label} =====\n{r['text']}")
        elif kind == "video":
            text_parts.append(f"\n===== UPLOADED VIDEO (reference only): {label} =====")

    return "\n".join(text_parts), images, tables


def stored_path(stored_as: str) -> Path:
    return UPLOAD_DIR / stored_as
