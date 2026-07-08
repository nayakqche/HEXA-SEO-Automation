"""
Renderer — converts the canonical blog JSON into Markdown, HTML, PDF, and DOCX.

The JSON is the source of truth (perfect for direct CMS import). The other
formats are generated from it so they always stay in sync. Paragraphs and
list items can carry an optional `links: [{anchor, href, kind}]` array — the
renderer turns those into real inline hyperlinks in every format.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement
from docx.shared import Inches, Pt, RGBColor
from xhtml2pdf import pisa


def _basename(src: str) -> str:
    return os.path.basename(src) if src else ""


def _image_exists(asset_dir: Path | None, src: str) -> bool:
    """True if we have the actual image file in asset_dir."""
    if not asset_dir or not src:
        return False
    return (asset_dir / _basename(src)).exists()


def _esc(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _table_data(block: dict) -> tuple[list[str], list[list[str]]]:
    """Normalise a table block into (headers, rows) of plain strings.

    Tolerates a missing `headers` (first row becomes the header) and pads short
    rows so the columns always line up.
    """
    headers = [str(h) for h in (block.get("headers") or [])]
    raw_rows = block.get("rows") or []
    rows = [[("" if c is None else str(c)) for c in r] for r in raw_rows]
    if not headers and rows:
        headers, rows = rows[0], rows[1:]
    width = max([len(headers)] + [len(r) for r in rows]) if (headers or rows) else 0
    headers = headers + [""] * (width - len(headers))
    rows = [r + [""] * (width - len(r)) for r in rows]
    return headers, rows


def _link_segments(text: str, links: list[dict]):
    """
    Split `text` into a sequence of (segment_text, link_or_none) tuples,
    splitting on anchor substrings (first match each). Order links by their
    first occurrence so substitutions don't overlap.
    """
    if not links:
        return [(text, None)]
    # Find first occurrence of each anchor; drop anchors not present.
    indexed = []
    for link in links:
        anchor = link.get("anchor", "")
        if not anchor:
            continue
        idx = text.find(anchor)
        if idx == -1:
            continue
        indexed.append((idx, anchor, link))
    indexed.sort(key=lambda t: t[0])

    segments: list[tuple[str, dict | None]] = []
    cursor = 0
    used = set()
    for idx, anchor, link in indexed:
        # Re-find from cursor so we never re-substitute already-consumed text.
        find_idx = text.find(anchor, cursor)
        if find_idx == -1 or id(link) in used:
            continue
        if find_idx > cursor:
            segments.append((text[cursor:find_idx], None))
        segments.append((anchor, link))
        cursor = find_idx + len(anchor)
        used.add(id(link))
    if cursor < len(text):
        segments.append((text[cursor:], None))
    return segments or [(text, None)]


# ── Markdown ───────────────────────────────────────────────────────────────

def _md_with_links(text: str, links: list[dict]) -> str:
    out = []
    for seg, link in _link_segments(text, links or []):
        if link:
            out.append(f"[{seg}]({link['href']})")
        else:
            out.append(seg)
    return "".join(out)


def render_markdown(post: dict) -> str:
    meta = post.get("meta", {})
    seo = post.get("seo", {})
    hero = post.get("hero", {}).get("image", {})

    lines: list[str] = ["---"]
    lines.append(f'title: "{seo.get("title", meta.get("title", ""))}"')
    lines.append(f'meta_description: "{seo.get("description", "")}"')
    lines.append(f'slug: "{post.get("slug", "")}"')
    if meta.get("category"): lines.append(f'category: "{meta["category"]}"')
    if meta.get("publishedDate"): lines.append(f'date: "{meta["publishedDate"]}"')
    if meta.get("author"): lines.append(f'author: "{meta["author"]}"')
    if meta.get("tags"):
        lines.append("tags: [" + ", ".join(f'"{t}"' for t in meta["tags"]) + "]")
    if seo.get("keywords"):
        lines.append("keywords: [" + ", ".join(f'"{k}"' for k in seo["keywords"]) + "]")
    lines.append("---\n")

    if hero.get("src"):
        lines.append(f'![{hero.get("alt", "")}]({_basename(hero["src"])})\n')

    title = meta.get("title") or seo.get("title", "")
    if title:
        lines.append(f"# {title}")
    if meta.get("subtitle"):
        lines.append(f"\n_{meta['subtitle']}_")
    lines.append("")

    for block in post.get("content", []):
        t = block.get("type")
        links = block.get("links") or []
        if t == "heading":
            level = max(2, min(6, int(block.get("level", 2))))
            lines.append("#" * level + " " + block.get("text", ""))
            lines.append("")
        elif t == "paragraph":
            lines.append(_md_with_links(block.get("text", ""), links))
            lines.append("")
        elif t == "list":
            marker = (lambda i: f"{i + 1}.") if block.get("style") == "ordered" else (lambda i: "-")
            for i, item in enumerate(block.get("items", [])):
                lines.append(f"{marker(i)} {_md_with_links(item, links)}")
            lines.append("")
        elif t == "image":
            lines.append(f"![{block.get('alt', '')}]({_basename(block.get('src', ''))})")
            if block.get("caption"):
                lines.append(f"\n*{block['caption']}*")
            lines.append("")
        elif t == "table":
            headers, rows = _table_data(block)
            if headers:
                lines.append("| " + " | ".join(headers) + " |")
                lines.append("| " + " | ".join("---" for _ in headers) + " |")
            for row in rows:
                lines.append("| " + " | ".join(row) + " |")
            if block.get("caption"):
                lines.append(f"\n*{block['caption']}*")
            lines.append("")

    return "\n".join(lines).rstrip() + "\n"


# ── HTML ───────────────────────────────────────────────────────────────────

_HTML_CSS = """
  body{font-family:-apple-system,Segoe UI,Roboto,Helvetica,Arial,sans-serif;
       max-width:780px;margin:0 auto;padding:2rem 1.25rem;color:#0f172a;line-height:1.7}
  .hero{width:100%;border-radius:14px;margin-bottom:1.25rem}
  .hero-placeholder{margin-bottom:1.25rem}
  .subtitle{color:#475569;font-size:1.1rem;margin-top:-0.5rem}
  .byline{color:#64748b;font-size:.88rem;margin-bottom:1.25rem}
  h1{font-size:2rem;line-height:1.2;margin-bottom:.25rem}
  h2{margin-top:2rem;color:#1d4ed8} h3{margin-top:1.4rem;color:#1e40af}
  a{color:#1d4ed8;text-decoration:underline}
  a.internal{color:#1e40af;font-weight:600;text-decoration:none;border-bottom:2px solid #c1d0f7}
  ul,ol{padding-left:1.4rem} li{margin:.35rem 0}
  figure{margin:1.5rem 0} figure img{width:100%;border-radius:12px}
  figcaption{color:#475569;font-size:.88rem;margin-top:.4rem;text-align:center}
  .image-placeholder{margin:1.5rem 0;padding:2rem 1.25rem;border:2px dashed #c1d0f7;
                     border-radius:12px;background:#f4f7ff;text-align:center}
  .image-placeholder .ip-label{font-size:.72rem;font-weight:700;color:#1d4ed8;
                               text-transform:uppercase}
  .image-placeholder .ip-alt{color:#1e293b;margin:.4rem 0 .15rem;font-size:.95rem}
  .image-placeholder .ip-cap{color:#475569;font-size:.85rem;font-style:italic}
  .tags{margin-top:2rem;display:flex;flex-wrap:wrap;gap:.4rem}
  .tag{font-size:.74rem;background:#e5edff;color:#1d4ed8;padding:.16rem .55rem;border-radius:999px}
  .table-wrap{overflow-x:auto}
  table{border-collapse:collapse;width:100%;margin:.25rem 0;font-size:.92rem}
  th,td{border:1px solid #dbe4f5;padding:.5rem .7rem;text-align:left;vertical-align:top}
  thead th{background:#eef2fb;color:#1e40af;font-weight:700}
  tbody tr:nth-child(even){background:#f7f9fe}
"""


def _image_placeholder_html(alt: str, caption: str = "", *, hero: bool = False) -> str:
    label = "Hero image" if hero else "Image"
    alt_html = f'<p class="ip-alt">{_esc(alt)}</p>' if alt else ""
    cap_html = f'<p class="ip-cap">{_esc(caption)}</p>' if caption else ""
    return (
        f'<div class="image-placeholder{" hero-placeholder" if hero else ""}">'
        f'<div class="ip-label">📷 {label} suggestion</div>'
        f'{alt_html}{cap_html}</div>'
    )


def _html_with_links(text: str, links: list[dict]) -> str:
    parts: list[str] = []
    for seg, link in _link_segments(text, links or []):
        if link:
            cls = "internal" if link.get("kind") == "internal" else "citation"
            target = "" if cls == "internal" else ' target="_blank" rel="noopener"'
            parts.append(
                f'<a href="{_esc(link["href"])}" class="{cls}"{target}>{_esc(seg)}</a>'
            )
        else:
            parts.append(_esc(seg))
    return "".join(parts)


def _html_blocks(blocks: Iterable[dict], asset_dir: Path | None) -> list[str]:
    out: list[str] = []
    for b in blocks:
        t = b.get("type")
        links = b.get("links") or []
        if t == "heading":
            lvl = max(2, min(6, int(b.get("level", 2))))
            out.append(f"<h{lvl}>{_esc(b.get('text', ''))}</h{lvl}>")
        elif t == "paragraph":
            out.append(f"<p>{_html_with_links(b.get('text', ''), links)}</p>")
        elif t == "list":
            tag = "ol" if b.get("style") == "ordered" else "ul"
            items = "".join(
                f"<li>{_html_with_links(it, links)}</li>" for it in b.get("items", [])
            )
            out.append(f"<{tag}>{items}</{tag}>")
        elif t == "image":
            src = _basename(b.get("src", ""))
            alt = b.get("alt", "")
            cap = b.get("caption", "")
            if src and _image_exists(asset_dir, src):
                cap_html = f"<figcaption>{_esc(cap)}</figcaption>" if cap else ""
                out.append(f'<figure><img src="{src}" alt="{_esc(alt)}">{cap_html}</figure>')
            else:
                out.append(_image_placeholder_html(alt, cap))
        elif t == "table":
            headers, rows = _table_data(b)
            thead = ("<thead><tr>" + "".join(f"<th>{_esc(h)}</th>" for h in headers)
                     + "</tr></thead>") if headers else ""
            tbody = "<tbody>" + "".join(
                "<tr>" + "".join(f"<td>{_esc(c)}</td>" for c in row) + "</tr>"
                for row in rows
            ) + "</tbody>"
            cap_html = f"<figcaption>{_esc(b.get('caption', ''))}</figcaption>" if b.get("caption") else ""
            out.append(f'<figure class="table-wrap"><table>{thead}{tbody}</table>{cap_html}</figure>')
    return out


def render_html(post: dict, asset_dir: Path | None = None) -> str:
    meta = post.get("meta", {})
    seo = post.get("seo", {})
    hero = post.get("hero", {}).get("image", {})

    title = _esc(meta.get("title") or seo.get("title", ""))
    subtitle = _esc(meta.get("subtitle", ""))
    desc = _esc(seo.get("description", ""))
    hero_src = hero.get("src", "")
    if hero_src and _image_exists(asset_dir, hero_src):
        hero_html = (
            f'<img class="hero" src="{_basename(hero_src)}" alt="{_esc(hero.get("alt", ""))}">'
        )
    elif hero_src or hero.get("alt"):
        hero_html = _image_placeholder_html(hero.get("alt", ""), hero=True)
    else:
        hero_html = ""
    byline_bits = []
    if meta.get("author"): byline_bits.append(_esc(meta["author"]))
    if meta.get("publishedDate"): byline_bits.append(_esc(meta["publishedDate"]))
    if meta.get("readTimeMinutes"): byline_bits.append(f"{meta['readTimeMinutes']} min read")
    byline = (' <span class="byline">' + " · ".join(byline_bits) + "</span>") if byline_bits else ""

    body = "\n".join(_html_blocks(post.get("content", []), asset_dir))
    tags_html = ""
    if meta.get("tags"):
        chips = "".join(f'<span class="tag">{_esc(t)}</span>' for t in meta["tags"])
        tags_html = f'<div class="tags">{chips}</div>'

    return f"""<!doctype html>
<html lang="en"><head>
<meta charset="utf-8"><meta name="viewport" content="width=device-width,initial-scale=1">
<title>{title}</title>
<meta name="description" content="{desc}">
<style>{_HTML_CSS}</style></head>
<body>
{hero_html}
<h1>{title}</h1>
{f'<p class="subtitle">{subtitle}</p>' if subtitle else ''}
{byline}
{body}
{tags_html}
</body></html>"""


# ── PDF (via xhtml2pdf) ────────────────────────────────────────────────────

def render_pdf(post: dict, output_path: Path, asset_dir: Path) -> None:
    """PDF render — uses the same HTML pipeline so missing images become placeholders."""
    html = render_html(post, asset_dir=asset_dir)

    def link_callback(uri: str, rel: str) -> str:
        if uri.startswith(("http://", "https://", "data:")):
            return uri
        p = (asset_dir / uri).resolve()
        return str(p) if p.exists() else uri

    with open(output_path, "wb") as fh:
        pisa.CreatePDF(src=html, dest=fh, link_callback=link_callback)


# ── DOCX (via python-docx) ─────────────────────────────────────────────────

_BLUE = RGBColor(0x1D, 0x4E, 0xD8)
_NAVY = RGBColor(0x1E, 0x40, 0xAF)
_SLATE = RGBColor(0x47, 0x55, 0x69)


def _add_text_run(paragraph, text: str, *, size: int = 11) -> None:
    if not text:
        return
    r = paragraph.add_run(text)
    r.font.size = Pt(size)


def _add_hyperlink(paragraph, url: str, text: str, *, color: RGBColor = _BLUE, size: int = 11) -> None:
    """Append a real Word hyperlink to a paragraph (python-docx has no native helper)."""
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), f"{color[0]:02X}{color[1]:02X}{color[2]:02X}")
    rPr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))  # half-points
    rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def _write_paragraph_with_links(paragraph, text: str, links: list[dict]) -> None:
    for seg, link in _link_segments(text, links or []):
        if link:
            _add_hyperlink(paragraph, link["href"], seg)
        else:
            _add_text_run(paragraph, seg)


def render_docx(post: dict, output_path: Path, asset_dir: Path) -> None:
    meta = post.get("meta", {})
    seo = post.get("seo", {})
    hero = post.get("hero", {}).get("image", {})
    doc = Document()

    title = meta.get("title") or seo.get("title", "")
    t = doc.add_paragraph()
    tr = t.add_run(title); tr.bold = True
    tr.font.size = Pt(22); tr.font.color.rgb = _NAVY

    if meta.get("subtitle"):
        sp = doc.add_paragraph()
        sr = sp.add_run(meta["subtitle"]); sr.italic = True
        sr.font.size = Pt(12); sr.font.color.rgb = _SLATE

    byline_bits = []
    if meta.get("author"): byline_bits.append(meta["author"])
    if meta.get("publishedDate"): byline_bits.append(meta["publishedDate"])
    if meta.get("readTimeMinutes"): byline_bits.append(f"{meta['readTimeMinutes']} min read")
    if byline_bits:
        bp = doc.add_paragraph()
        br = bp.add_run(" · ".join(byline_bits))
        br.font.size = Pt(9); br.font.color.rgb = _SLATE

    hero_path = asset_dir / _basename(hero.get("src", "")) if hero.get("src") else None
    if hero_path and hero_path.exists():
        doc.add_picture(str(hero_path), width=Inches(6))
    elif hero.get("alt"):
        hp = doc.add_paragraph()
        hr = hp.add_run(f"[ Hero image: {hero['alt']} ]")
        hr.italic = True; hr.font.size = Pt(9); hr.font.color.rgb = _SLATE

    for block in post.get("content", []):
        bt = block.get("type")
        links = block.get("links") or []
        if bt == "heading":
            lvl = max(1, min(4, int(block.get("level", 2))))
            doc.add_heading(block.get("text", ""), level=lvl)
        elif bt == "paragraph":
            p = doc.add_paragraph()
            _write_paragraph_with_links(p, block.get("text", ""), links)
        elif bt == "list":
            style = "List Number" if block.get("style") == "ordered" else "List Bullet"
            for item in block.get("items", []):
                p = doc.add_paragraph(style=style)
                _write_paragraph_with_links(p, item, links)
        elif bt == "image":
            img_path = asset_dir / _basename(block.get("src", ""))
            if img_path.exists():
                doc.add_picture(str(img_path), width=Inches(6))
                if block.get("caption"):
                    cp = doc.add_paragraph()
                    cr = cp.add_run(block["caption"]); cr.italic = True
                    cr.font.size = Pt(9); cr.font.color.rgb = _SLATE
            elif block.get("alt") or block.get("caption"):
                # No image file (Pexels skipped/failed) → leave a clear marker.
                ip = doc.add_paragraph()
                ir = ip.add_run(f"[ Image: {block.get('alt', '')} ]")
                ir.italic = True; ir.font.size = Pt(9); ir.font.color.rgb = _SLATE
                if block.get("caption"):
                    cp = doc.add_paragraph()
                    cr = cp.add_run(block["caption"]); cr.italic = True
                    cr.font.size = Pt(9); cr.font.color.rgb = _SLATE
        elif bt == "table":
            headers, rows = _table_data(block)
            all_rows = ([headers] if headers else []) + rows
            if all_rows:
                table = doc.add_table(rows=len(all_rows), cols=len(all_rows[0]))
                table.style = "Light Grid Accent 1"
                for ri, row in enumerate(all_rows):
                    for ci, val in enumerate(row):
                        cell = table.rows[ri].cells[ci]
                        cell.text = val
                        if ri == 0 and headers:
                            for run in cell.paragraphs[0].runs:
                                run.bold = True
                if block.get("caption"):
                    cp = doc.add_paragraph()
                    cr = cp.add_run(block["caption"]); cr.italic = True
                    cr.font.size = Pt(9); cr.font.color.rgb = _SLATE

    if meta.get("tags"):
        doc.add_paragraph()
        fp = doc.add_paragraph()
        fr = fp.add_run("Tags: " + ", ".join(meta["tags"]))
        fr.font.size = Pt(9); fr.font.color.rgb = _SLATE

    doc.save(str(output_path))
